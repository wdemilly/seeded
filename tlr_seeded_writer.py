import streamlit as st
import anthropic
import re
import math
import io
from docx import Document
from docx.shared import Pt

# ─── PAGE CONFIG ─────────────────────────────────────────────────
st.set_page_config(page_title="TLR Seeded Outline Writer", layout="wide")
st.title("TLR — Seeded Outline Writer")

# ─── SIDEBAR ─────────────────────────────────────────────────────
with st.sidebar:
    st.header("Configuration")
    api_key = st.text_input("Anthropic API Key", type="password")

    st.subheader("Writing Model")
    writing_model = st.selectbox("Model", [
        "claude-sonnet-4-20250514",
        "claude-opus-4-20250514",
        "claude-haiku-4-5-20251001",
    ], index=0)
    writing_temp = st.slider("Writing Temperature", 0.0, 1.0, 1.0, 0.1)

    st.subheader("Revision")
    auto_revise = st.checkbox("Auto-Revise", value=False)
    if auto_revise:
        revision_model = st.selectbox("Revision Model", [
            "claude-sonnet-4-20250514",
            "claude-opus-4-20250514",
            "claude-haiku-4-5-20251001",
        ], index=0, key="rev_model")
        revision_temp = st.slider("Revision Temperature", 0.0, 1.0, 0.3, 0.1)
        revision_passes = st.slider("Revision Passes", 1, 3, 2)

    st.divider()
    st.subheader("Upload Documents")
    outline_file = st.file_uploader("Chapter Outline (seeded)", type=["docx", "txt"])
    source_file = st.file_uploader("Combined Source Texts", type=["docx", "txt"])
    profiles_file = st.file_uploader("Character Profiles", type=["docx", "txt"])

    st.divider()
    st.subheader("Writing Prompt")
    default_prompt = """You are the author of the combined source texts. You wrote all of them. The character profiles are your notes. The outline is your plan for this chapter.

Write the chapter now, the way you wrote the source texts. One continuous pass, first sentence to last. Do not draft short and expand."""
    writing_prompt = st.text_area("Prompt", value=default_prompt, height=180)


# ─── FILE READERS ────────────────────────────────────────────────
def read_upload(uploaded_file):
    """Read an uploaded file and return its text."""
    if uploaded_file is None:
        return None
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="replace")
    elif name.endswith(".docx"):
        doc = Document(io.BytesIO(uploaded_file.read()))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        return uploaded_file.read().decode("utf-8", errors="replace")


# ─── API CALL ────────────────────────────────────────────────────
def call_api(client, model, temperature, messages, max_tokens=8192):
    """Make a single API call and return the text response."""
    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        messages=messages,
    )
    text = response.content[0].text
    usage = response.usage
    return text, usage


# ─── SCORER ──────────────────────────────────────────────────────
def score_chapter(text):
    """Run detection metrics. Returns dict of results."""
    words = text.split()
    word_count = len(words)
    if word_count == 0:
        return None
    kw = word_count / 1000

    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z"\u201C])', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    slens = [len(s.split()) for s in sentences]

    if len(slens) < 2:
        mean_len = sum(slens) / max(len(slens), 1)
        cv = 0
    else:
        mean_len = sum(slens) / len(slens)
        variance = sum((x - mean_len) ** 2 for x in slens) / len(slens)
        cv = math.sqrt(variance) / mean_len if mean_len > 0 else 0

    counts = {
        "Em dashes (—)": len(re.findall(r'—', text)),
        "As though / as if": len(re.findall(r'as though|as if', text, re.I)),
        "The way [pronoun]": len(re.findall(r'the way (?:he|she|I|they|it|men|people|women|a man|a woman)\b', text, re.I)),
        "Of someone/person who": len(re.findall(r'of (?:someone|a man|a woman|a person|people|a creature) who', text, re.I)),
        "With the [adj] of": len(re.findall(r'with the \w+ (?:\w+ )?of', text, re.I)),
        "Negation-leading": len(re.findall(r'\b(?:It|That|This|He|She|I) (?:was|had|did|could) not\b', text)),
        "Not X but Y": len(re.findall(r'not \w+[\w, ]* but ', text, re.I)),
        "Which meant/was": len(re.findall(r'which (?:meant|was|had|would|could|did)', text, re.I)),
        "Editorial (This was what/how)": len(re.findall(r'(?:This|That) was (?:what|how|the)', text, re.I)),
        "The kind/sort of": len(re.findall(r'the kind of|the sort of', text, re.I)),
        "Semicolons": len(re.findall(r';', text)),
        "Metacognitive verbs": len(re.findall(r'\b(?:noticed|noted|filed|registered|understood|recognised|recognized|catalogued|assessed|calculated)\b', text, re.I)),
    }

    thresholds = {
        "Em dashes (—)": (1.0, 2.5),
        "As though / as if": (0.2, 0.6),
        "The way [pronoun]": (0.5, 1.0),
        "Of someone/person who": (0.3, 0.8),
        "With the [adj] of": (0.3, 0.8),
        "Negation-leading": (2.0, 4.0),
        "Not X but Y": (0.3, 0.8),
        "Which meant/was": (0.5, 1.5),
        "Editorial (This was what/how)": (0.2, 0.5),
        "The kind/sort of": (0.2, 0.5),
        "Semicolons": (0.6, 1.2),
        "Metacognitive verbs": (0.3, 1.0),
    }

    metrics = []
    for name, count in counts.items():
        per_kw = count / kw if kw > 0 else 0
        g, y = thresholds[name]
        if per_kw <= g:
            status = "GREEN"
        elif per_kw <= y:
            status = "YELLOW"
        else:
            status = "RED"
        metrics.append({"name": name, "count": count, "per_kw": round(per_kw, 2), "status": status})

    # Ratio metrics
    periods = text.count('.')
    commas = text.count(',')
    p2c = periods / commas if commas > 0 else 99
    p2c_status = "GREEN" if p2c >= 1.8 else ("YELLOW" if p2c >= 1.5 else "RED")

    cv_status = "GREEN" if cv >= 1.2 else ("YELLOW" if cv >= 1.0 else "RED")

    dialogue_matches = re.findall(r'[\u201C"][^\u201D"]*[\u201D"]', text)
    dialogue_words = sum(len(m.split()) for m in dialogue_matches)
    dial_pct = 100 * dialogue_words / word_count
    dial_status = "GREEN" if dial_pct >= 15 else ("YELLOW" if dial_pct >= 8 else "RED")

    metrics.append({"name": "Period:comma ratio", "count": round(p2c, 3), "per_kw": None, "status": p2c_status})
    metrics.append({"name": "Sentence length CV", "count": round(cv, 3), "per_kw": None, "status": cv_status})
    metrics.append({"name": "Dialogue density %", "count": round(dial_pct, 1), "per_kw": None, "status": dial_status})

    red = sum(1 for m in metrics if m["status"] == "RED")
    green = sum(1 for m in metrics if m["status"] == "GREEN")
    yellow = len(metrics) - red - green

    if red >= 4:
        risk = "HIGH RISK"
    elif red >= 2:
        risk = "MODERATE RISK"
    else:
        risk = "LOW RISK"

    return {
        "metrics": metrics,
        "risk": risk,
        "red": red,
        "green": green,
        "yellow": yellow,
        "word_count": word_count,
        "sentences": len(sentences),
        "mean_len": round(mean_len, 1),
    }


def display_scorecard(score, label=""):
    """Display the scorecard in Streamlit."""
    if score is None:
        st.warning("No text to score.")
        return

    risk_color = {"HIGH RISK": "red", "MODERATE RISK": "orange", "LOW RISK": "green"}
    col = risk_color.get(score["risk"], "grey")
    st.markdown(f"### {label} — :{col}[{score['risk']}]")
    st.markdown(f"**{score['green']} GREEN** · **{score['yellow']} YELLOW** · **{score['red']} RED** · {score['word_count']} words · {score['sentences']} sentences · mean {score['mean_len']} words/sentence")

    for m in score["metrics"]:
        icon = {"GREEN": "🟢", "YELLOW": "🟡", "RED": "🔴"}[m["status"]]
        if m["per_kw"] is not None:
            st.text(f"  {icon} {m['name']:<32} {m['count']:>3}  ({m['per_kw']}/1kw)")
        else:
            st.text(f"  {icon} {m['name']:<32} {m['count']}")


def make_docx(text):
    """Create a .docx file from text and return bytes."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for para in text.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── REVISION PROMPT ─────────────────────────────────────────────
REVISION_PROMPT = """You are a prose editor. The chapter below has specific patterns that AI detection systems flag. Find and fix these WITHOUT rewriting good prose:

1. Replace any em dashes (—) with periods or commas
2. Delete any "as though" or "as if" + interpretation (keep the action, cut the simile)
3. Delete any "with the [adjective] [noun] of someone who" constructions (keep the action)
4. Delete any "the way [pronoun] [verb] when" constructions
5. Break any sentence over 35 words into two sentences
6. Delete editorial commentary: "This was what/how..." sentences
7. Delete "which meant" / "which was" chains — break into separate sentences or cut
8. Where you delete a construction, close the gap — do not generate a replacement image or metaphor

Preserve dialogue exactly. Preserve voice. Cut, don't rewrite. The chapter will be shorter. That is correct.

Return only the revised chapter text, no commentary.

CHAPTER:
"""


# ─── MAIN ────────────────────────────────────────────────────────
if st.button("Write Chapter", type="primary", disabled=not api_key):
    if not all([outline_file, source_file, profiles_file]):
        st.error("Upload all three documents: outline, source texts, and character profiles.")
        st.stop()

    outline_text = read_upload(outline_file)
    source_text = read_upload(source_file)
    profiles_text = read_upload(profiles_file)

    if not all([outline_text, source_text, profiles_text]):
        st.error("Could not read one or more uploaded files.")
        st.stop()

    st.info(f"Inputs loaded — Outline: {len(outline_text.split())} words · Source: {len(source_text.split())} words · Profiles: {len(profiles_text.split())} words")

    client = anthropic.Anthropic(api_key=api_key)

    # ── WRITE ────────────────────────────────────────────────
    user_content = f"""=== CHARACTER PROFILES ===
{profiles_text}

=== SOURCE TEXTS ===
{source_text}

=== CHAPTER OUTLINE ===
{outline_text}

=== INSTRUCTION ===
{writing_prompt}"""

    with st.spinner(f"Writing chapter with {writing_model} at temp {writing_temp}..."):
        try:
            chapter_text, usage = call_api(
                client, writing_model, writing_temp,
                [{"role": "user", "content": user_content}]
            )
        except Exception as e:
            st.error(f"API error: {e}")
            st.stop()

    st.success(f"Chapter written — {len(chapter_text.split())} words · {usage.input_tokens} input tokens · {usage.output_tokens} output tokens")

    # ── SCORE ORIGINAL ───────────────────────────────────────
    original_score = score_chapter(chapter_text)
    display_scorecard(original_score, "Original")

    st.subheader("Original Chapter")
    st.text_area("", chapter_text, height=400, key="original_text")
    st.download_button(
        "Download Original (.docx)",
        make_docx(chapter_text),
        file_name="chapter_original.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # ── REVISE (optional) ────────────────────────────────────
    if auto_revise:
        current_text = chapter_text
        for pass_num in range(1, revision_passes + 1):
            with st.spinner(f"Revision pass {pass_num} with {revision_model} at temp {revision_temp}..."):
                try:
                    revised_text, rev_usage = call_api(
                        client, revision_model, revision_temp,
                        [{"role": "user", "content": REVISION_PROMPT + current_text}]
                    )
                except Exception as e:
                    st.error(f"Revision pass {pass_num} error: {e}")
                    break

            current_text = revised_text
            rev_score = score_chapter(revised_text)
            display_scorecard(rev_score, f"Revision {pass_num}")

            st.subheader(f"Revision {pass_num}")
            st.text_area("", revised_text, height=400, key=f"rev_{pass_num}_text")
            st.download_button(
                f"Download Rev {pass_num} (.docx)",
                make_docx(revised_text),
                file_name=f"chapter_rev{pass_num}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_rev_{pass_num}",
            )

elif not api_key:
    st.warning("Enter your Anthropic API key in the sidebar to begin.")
