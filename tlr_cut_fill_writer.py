import streamlit as st
import anthropic
import re
import math
import io
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="TLR Cut & Fill Writer", layout="wide")
st.title("TLR — Cut & Fill Writer")

# ─── SESSION STATE ───────────────────────────────────────────────
for key in ["chapters", "scores", "run_count", "cache_stats"]:
    if key not in st.session_state:
        st.session_state[key] = {} if key in ("chapters", "scores") else (0 if key == "run_count" else [])

# ─── SIDEBAR ─────────────────────────────────────────────────────
with st.sidebar:
    st.header("Configuration")
    api_key = st.text_input("Anthropic API Key", type="password")

    st.subheader("Writing Model")
    writing_model = st.selectbox("Model", [
        "claude-sonnet-4-20250514",
        "claude-opus-4-20250514",
        "claude-haiku-4-5-20251001",
    ], index=0)
    writing_temp = st.slider("Writing Temperature", 0.0, 1.0, 1.0, 0.1)

    st.subheader("Pipeline")
    run_cuts = st.checkbox("Run Cuts Pass", value=True)
    run_fill = st.checkbox("Run Mechanical Fill", value=False,
                           help="After cuts, fill gaps using closed-vocabulary Oulipo constraints")
    if run_fill:
        fill_model = st.selectbox("Fill Model", [
            "claude-sonnet-4-20250514",
            "claude-opus-4-20250514",
            "claude-haiku-4-5-20251001",
        ], index=0, key="fill_model")
        fill_temp = st.slider("Fill Temperature", 0.0, 1.0, 0.7, 0.1)

    st.divider()
    st.subheader("Upload Documents")
    outline_file = st.file_uploader("Chapter Outline (seeded)", type=["docx", "txt"])
    source_file = st.file_uploader("Combined Source Texts", type=["docx", "txt"])
    profiles_file = st.file_uploader("Character Profiles", type=["docx", "txt"])

    st.divider()
    st.subheader("Writing Prompt")
    default_prompt = """You are the author of the combined source texts. You wrote all of them. The character profiles are your notes. The outline is your plan for this chapter.

Write the chapter now, the way you wrote the source texts. One continuous pass, first sentence to last. Do not draft short and expand. Write at full length — every beat in the outline gets its full weight. Aim for a chapter that is substantially longer than you think it needs to be. Dwell in scenes. Let dialogue run. Add physical business between spoken lines. Do not summarise what the outline says to dramatise."""
    writing_prompt = st.text_area("Prompt", value=default_prompt, height=200)


def read_upload(uploaded_file):
    if uploaded_file is None:
        return None
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="replace")
    elif name.endswith(".docx"):
        doc = Document(io.BytesIO(uploaded_file.read()))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        return uploaded_file.read().decode("utf-8", errors="replace")


# ─── API CALLS ───────────────────────────────────────────────────

def call_api_cached(client, model, temperature, source_text, profiles_text,
                    outline_text, prompt_text, max_tokens=8192):
    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        system=[{
            "type": "text",
            "text": f"=== CHARACTER PROFILES ===\n{profiles_text}\n\n=== SOURCE TEXTS ===\n{source_text}",
            "cache_control": {"type": "ephemeral"},
        }],
        messages=[{
            "role": "user",
            "content": f"=== CHAPTER OUTLINE ===\n{outline_text}\n\n=== INSTRUCTION ===\n{prompt_text}",
        }],
    )
    text = response.content[0].text
    u = response.usage
    return text, {
        "input": u.input_tokens, "output": u.output_tokens,
        "cache_creation": getattr(u, "cache_creation_input_tokens", 0),
        "cache_read": getattr(u, "cache_read_input_tokens", 0),
    }


def call_api_plain(client, model, temperature, prompt, max_tokens=8192):
    response = client.messages.create(
        model=model, max_tokens=max_tokens, temperature=temperature,
        messages=[{"role": "user", "content": prompt}],
    )
    text = response.content[0].text
    u = response.usage
    return text, {
        "input": u.input_tokens, "output": u.output_tokens,
        "cache_creation": getattr(u, "cache_creation_input_tokens", 0),
        "cache_read": getattr(u, "cache_read_input_tokens", 0),
    }


# ─── CUTS PROMPT ─────────────────────────────────────────────────

CUTS_PROMPT = """You are a mechanical editor. You DELETE. You do NOT rewrite. You do NOT generate replacement text.

Read the chapter below. Apply these cuts:

1. DELETE any em dash (—) and the clause it introduces. Keep the text before the dash. End with a period.
2. DELETE any "as though" or "as if" + everything after it in that sentence. Keep the action before it. End with a period.
3. DELETE any "with the [adjective] [noun] of someone/a man/a woman who" construction. Keep the subject and verb before it. End with a period.
4. DELETE any "the way [pronoun] [verb] when" construction. Keep the action. End with a period.
5. DELETE any sentence beginning with "This was what" or "This was how" or "That was the" where the narrator explains what just happened.
6. DELETE any "which meant" / "which was" / "which had" trailing clause. End the sentence at the word before "which". Add a period.
7. DELETE any "the kind of" / "the sort of" construction and its trailing clause.
8. DELETE any sentence that interprets a scene the reader has already witnessed — editorial commentary.
9. DELETE any sentence containing "of someone who" / "of a man who" / "of a woman who" / "of a person who".

WHERE YOU DELETE: close the gap. Do not insert a replacement phrase, image, metaphor, or explanation. Shorter is correct.

PRESERVE: all dialogue exactly as written. All physical actions. All scene transitions. All plot beats.

The chapter WILL be shorter. That is correct. Do not try to maintain word count.

Return ONLY the cut chapter text. No commentary, no change log.

CHAPTER:
"""


# ─── FILL PROMPT BUILDER ─────────────────────────────────────────

def build_fill_prompt(cut_chapter, outline_text):
    """Build the mechanical fill prompt with Oulipo constraints."""

    # Extract the chapter's existing vocabulary (closed stock)
    words = re.findall(r"[a-zA-Z']+", cut_chapter.lower())
    word_freq = {}
    for w in words:
        word_freq[w] = word_freq.get(w, 0) + 1
    # Keep words that appear 2+ times (the chapter's real vocabulary)
    vocab = sorted([w for w, c in word_freq.items() if c >= 2])
    vocab_sample = ", ".join(vocab[:120])

    # Get sentence lengths from surviving text for rhythm targets
    sents = re.split(r'(?<=[.!?])\s+(?=[A-Z"\u201C])', cut_chapter)
    sents = [s.strip() for s in sents if s.strip()]
    slens = [len(s.split()) for s in sents]
    if slens:
        short_target = min(slens)
        long_target = max(slens)
        median_len = sorted(slens)[len(slens) // 2]
    else:
        short_target, long_target, median_len = 4, 30, 14

    prompt = f"""You are completing a partially written chapter. This is NOT a revision. You are NOT editing. You are writing missing sections of a draft that has gaps.

Below is a chapter draft with [CONTINUE HERE — beat: description] markers where material is needed. Also below is the chapter outline.

RULES — THESE ARE ABSOLUTE:

1. VOCABULARY CONSTRAINT: You may ONLY use words that already appear in the draft below. Here is the chapter's word stock: {vocab_sample}
   Function words (the, a, an, in, on, at, to, of, for, and, but, or, with, is, was, were, had, have, has, did, do, not, no, it, he, she, they, I, my, his, her, we, you, this, that, from, by, as, if, so, up, out, its, been, than, who, what, when, where, how, which, will, would, could, should, can, may, might, just, now, then, here, there, each, every, all, some, any, much, more, most, very, too, also, still, only, even, back, down, over, after, before, between, through, about) are always allowed.

2. SENTENCE LENGTH CONSTRAINT: Each new sentence must be within 3 words of the length of the sentence immediately BEFORE the [CONTINUE HERE] marker it follows. If the preceding sentence is 12 words, your sentence must be 9-15 words. If no preceding sentence exists, use {median_len} words ± 3.

3. You are writing NEW narration for the beats described in the markers. Use the outline to know what happens. Write concrete physical action and dialogue. No interpretation. No editorial commentary. No metaphors. No similes.

4. Write 2-4 sentences per marker. No more.

5. Do NOT touch any existing text. Output the COMPLETE chapter — all existing text preserved exactly, with your new sentences inserted at each marker.

=== CHAPTER OUTLINE ===
{outline_text}

=== CHAPTER DRAFT WITH MARKERS ===
{cut_chapter}

Output the complete chapter with markers replaced by your new sentences. No commentary."""

    return prompt


def insert_continuation_markers(cut_chapter, original_chapter, outline_text):
    """Compare cut vs original to find where material was lost,
    then insert [CONTINUE HERE] markers at those gaps."""

    cut_paras = [p.strip() for p in cut_chapter.split('\n') if p.strip()]
    orig_paras = [p.strip() for p in original_chapter.split('\n') if p.strip()]

    # Find paragraphs that got significantly shorter or where consecutive
    # sentences were removed (gaps in the flow)
    result_paras = []
    for i, para in enumerate(cut_paras):
        result_paras.append(para)

        # Check if this paragraph is much shorter than the corresponding original
        # or if there's a logical gap (end of one scene beat, start of another
        # with nothing between)
        cut_sents = re.split(r'(?<=[.!?])\s+', para)
        cut_words = len(para.split())

        # Find the best-matching original paragraph
        best_match = 0
        best_idx = -1
        for j, opara in enumerate(orig_paras):
            # Simple overlap check
            overlap = len(set(para.split()[:10]) & set(opara.split()[:10]))
            if overlap > best_match:
                best_match = overlap
                best_idx = j

        if best_idx >= 0:
            orig_words = len(orig_paras[best_idx].split())
            word_loss = orig_words - cut_words
            if word_loss > 40:
                # Significant material was cut from this section
                result_paras.append(
                    f"\n[CONTINUE HERE — this section lost ~{word_loss} words. "
                    f"Add physical action or concrete detail for this scene beat.]\n"
                )

    # Also check for overall word deficit and add markers at natural break points
    cut_total = len(cut_chapter.split())
    orig_total = len(original_chapter.split())
    deficit = orig_total - cut_total

    if deficit > 200:
        # Add a few more markers at paragraph boundaries where there's no dialogue
        non_dialogue_indices = []
        for i, para in enumerate(result_paras):
            if not re.search(r'[\u201C"][^\u201D"]*[\u201D"]', para) and len(para.split()) > 20:
                non_dialogue_indices.append(i)

        # Insert markers at every 3rd non-dialogue paragraph
        inserted = 0
        for idx in non_dialogue_indices[::3]:
            adj_idx = idx + inserted + 1
            if adj_idx < len(result_paras):
                result_paras.insert(adj_idx,
                    "\n[CONTINUE HERE — add 2-3 sentences of physical business or concrete observation.]\n"
                )
                inserted += 1
            if inserted >= 4:
                break

    return '\n\n'.join(result_paras)


# ─── SCORER ──────────────────────────────────────────────────────

def score_chapter(text):
    words = text.split()
    word_count = len(words)
    if word_count == 0:
        return None
    kw = word_count / 1000

    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z"\u201C])', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    slens = [len(s.split()) for s in sentences]

    if len(slens) < 2:
        mean_len = sum(slens) / max(len(slens), 1)
        cv = 0
    else:
        mean_len = sum(slens) / len(slens)
        variance = sum((x - mean_len) ** 2 for x in slens) / len(slens)
        cv = math.sqrt(variance) / mean_len if mean_len > 0 else 0

    counts = {
        "Em dashes": len(re.findall(r'\u2014', text)),
        "As though / as if": len(re.findall(r'as though|as if', text, re.I)),
        "The way [pronoun]": len(re.findall(r'the way (?:he|she|I|they|it|men|people|women|a man|a woman)\b', text, re.I)),
        "Of someone/person who": len(re.findall(r'of (?:someone|a man|a woman|a person|people|a creature) who', text, re.I)),
        "With the [adj] of": len(re.findall(r'with the \w+ (?:\w+ )?of', text, re.I)),
        "Negation-leading": len(re.findall(r'\b(?:It|That|This|He|She|I) (?:was|had|did|could) not\b', text)),
        "Not X but Y": len(re.findall(r'not \w+[\w, ]* but ', text, re.I)),
        "Which meant/was": len(re.findall(r'which (?:meant|was|had|would|could|did)', text, re.I)),
        "Editorial commentary": len(re.findall(r'(?:This|That) was (?:what|how|the)', text, re.I)),
        "The kind/sort of": len(re.findall(r'the kind of|the sort of', text, re.I)),
        "Semicolons": len(re.findall(r';', text)),
        "Metacognitive verbs": len(re.findall(r'\b(?:noticed|noted|filed|registered|understood|recognised|recognized|catalogued|assessed|calculated)\b', text, re.I)),
    }

    thresholds = {
        "Em dashes": (1.0, 2.5), "As though / as if": (0.2, 0.6),
        "The way [pronoun]": (0.5, 1.0), "Of someone/person who": (0.3, 0.8),
        "With the [adj] of": (0.3, 0.8), "Negation-leading": (2.0, 4.0),
        "Not X but Y": (0.3, 0.8), "Which meant/was": (0.5, 1.5),
        "Editorial commentary": (0.2, 0.5), "The kind/sort of": (0.2, 0.5),
        "Semicolons": (0.6, 1.2), "Metacognitive verbs": (0.3, 1.0),
    }

    metrics = []
    for name, count in counts.items():
        per_kw = count / kw if kw > 0 else 0
        g, y = thresholds[name]
        status = "GREEN" if per_kw <= g else ("YELLOW" if per_kw <= y else "RED")
        metrics.append({"name": name, "count": count, "per_kw": round(per_kw, 2), "status": status})

    periods = text.count('.')
    commas = text.count(',')
    p2c = periods / commas if commas > 0 else 99
    p2c_status = "GREEN" if p2c >= 1.8 else ("YELLOW" if p2c >= 1.5 else "RED")
    cv_status = "GREEN" if cv >= 1.2 else ("YELLOW" if cv >= 1.0 else "RED")

    dialogue_matches = re.findall(r'[\u201C"][^\u201D"]*[\u201D"]', text)
    dialogue_words = sum(len(m.split()) for m in dialogue_matches)
    dial_pct = 100 * dialogue_words / word_count
    dial_status = "GREEN" if dial_pct >= 15 else ("YELLOW" if dial_pct >= 8 else "RED")

    metrics.append({"name": "Period:comma ratio", "count": round(p2c, 3), "per_kw": None, "status": p2c_status})
    metrics.append({"name": "Sentence length CV", "count": round(cv, 3), "per_kw": None, "status": cv_status})
    metrics.append({"name": "Dialogue density %", "count": round(dial_pct, 1), "per_kw": None, "status": dial_status})

    red = sum(1 for m in metrics if m["status"] == "RED")
    green = sum(1 for m in metrics if m["status"] == "GREEN")
    yellow = len(metrics) - red - green

    risk = "HIGH RISK" if red >= 4 else ("MODERATE RISK" if red >= 2 else "LOW RISK")

    return {
        "metrics": metrics, "risk": risk,
        "red": red, "green": green, "yellow": yellow,
        "word_count": word_count, "sentences": len(sentences),
        "mean_len": round(mean_len, 1),
    }


def display_scorecard(score, label=""):
    if score is None:
        return
    risk_color = {"HIGH RISK": "red", "MODERATE RISK": "orange", "LOW RISK": "green"}
    col = risk_color.get(score["risk"], "grey")
    st.markdown(f"### {label} — :{col}[{score['risk']}]")
    st.markdown(f"**{score['green']} GREEN** · **{score['yellow']} YELLOW** · **{score['red']} RED** · {score['word_count']} words · {score['sentences']} sentences · mean {score['mean_len']} w/s")
    for m in score["metrics"]:
        icon = {"GREEN": "\u2705", "YELLOW": "\u26A0\uFE0F", "RED": "\u274C"}[m["status"]]
        if m["per_kw"] is not None:
            st.text(f"  {icon} {m['name']:<32} {m['count']:>3}  ({m['per_kw']}/1kw)")
        else:
            st.text(f"  {icon} {m['name']:<32} {m['count']}")


def make_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for para in text.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── DISPLAY STORED RESULTS ─────────────────────────────────────
if st.session_state.chapters:
    for key in sorted(st.session_state.chapters.keys()):
        label = key.replace("_", " ").title()
        score = st.session_state.scores.get(key)
        text = st.session_state.chapters[key]

        if score:
            display_scorecard(score, f"Run {st.session_state.run_count} — {label}")

        with st.expander(f"{label} — {len(text.split())} words", expanded=(key == "01_original")):
            st.text_area("", text, height=300, key=f"ta_{st.session_state.run_count}_{key}")
            st.download_button(
                f"Download {label} (.docx)",
                make_docx(text),
                file_name=f"chapter_{key}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{st.session_state.run_count}_{key}",
            )

    if st.session_state.cache_stats:
        with st.expander("Cache & Token Stats"):
            for stat in st.session_state.cache_stats:
                ci = stat.get("info", {})
                cache_read = ci.get("cache_read", 0)
                cache_write = ci.get("cache_creation", 0)
                hit = "CACHE HIT" if cache_read > 0 else ("CACHE WRITE" if cache_write > 0 else "NO CACHE")
                st.text(f"  {stat.get('label','')}: {ci.get('input',0)} in / {ci.get('output',0)} out / {cache_write} write / {cache_read} read — {hit}")

    st.divider()


# ─── WRITE BUTTON ────────────────────────────────────────────────
if st.button("Write Chapter", type="primary", disabled=not api_key):
    if not all([outline_file, source_file, profiles_file]):
        st.error("Upload all three documents: outline, source texts, and character profiles.")
        st.stop()

    outline_text = read_upload(outline_file)
    source_text = read_upload(source_file)
    profiles_text = read_upload(profiles_file)

    if not all([outline_text, source_text, profiles_text]):
        st.error("Could not read one or more uploaded files.")
        st.stop()

    st.info(f"Inputs — Outline: {len(outline_text.split())} w · Source: {len(source_text.split())} w · Profiles: {len(profiles_text.split())} w")

    client = anthropic.Anthropic(api_key=api_key)

    st.session_state.chapters = {}
    st.session_state.scores = {}
    st.session_state.cache_stats = []
    st.session_state.run_count += 1

    # ── PASS 1: WRITE (extended length) ──────────────────────────
    with st.spinner(f"Pass 1: Writing with {writing_model} at temp {writing_temp}..."):
        try:
            chapter_text, cache_info = call_api_cached(
                client, writing_model, writing_temp,
                source_text, profiles_text, outline_text, writing_prompt,
            )
        except Exception as e:
            st.error(f"API error: {e}")
            st.stop()

    st.session_state.chapters["01_original"] = chapter_text
    st.session_state.scores["01_original"] = score_chapter(chapter_text)
    st.session_state.cache_stats.append({"label": "Write", "info": cache_info})

    # ── PASS 2: CUTS ONLY ────────────────────────────────────────
    if run_cuts:
        with st.spinner("Pass 2: Cuts only (delete, no rewrite)..."):
            try:
                cut_text, cut_info = call_api_plain(
                    client, writing_model, 0.3,
                    CUTS_PROMPT + chapter_text,
                )
            except Exception as e:
                st.error(f"Cuts pass error: {e}")
                cut_text = None

        if cut_text:
            st.session_state.chapters["02_cuts"] = cut_text
            st.session_state.scores["02_cuts"] = score_chapter(cut_text)
            st.session_state.cache_stats.append({"label": "Cuts", "info": cut_info})

            orig_wc = len(chapter_text.split())
            cut_wc = len(cut_text.split())
            st.info(f"Cuts: {orig_wc} → {cut_wc} words ({orig_wc - cut_wc} removed, {100*cut_wc/orig_wc:.0f}% retained)")

            # ── PASS 3: MECHANICAL FILL ──────────────────────────
            if run_fill and cut_text:
                with st.spinner("Pass 3: Mechanical fill (Oulipo constraints)..."):
                    # Insert continuation markers where material was lost
                    marked_chapter = insert_continuation_markers(
                        cut_text, chapter_text, outline_text
                    )

                    # Check if any markers were actually inserted
                    marker_count = marked_chapter.count("[CONTINUE HERE")
                    if marker_count == 0:
                        st.warning("No significant gaps detected after cuts. Skipping fill pass.")
                    else:
                        st.info(f"Fill pass: {marker_count} continuation markers inserted")

                        fill_prompt = build_fill_prompt(marked_chapter, outline_text)

                        try:
                            filled_text, fill_info = call_api_plain(
                                client, fill_model, fill_temp,
                                fill_prompt,
                            )
                        except Exception as e:
                            st.error(f"Fill pass error: {e}")
                            filled_text = None

                        if filled_text:
                            # Clean any remaining markers the model didn't fill
                            filled_text = re.sub(r'\[CONTINUE HERE[^\]]*\]', '', filled_text).strip()
                            filled_text = re.sub(r'\n{3,}', '\n\n', filled_text)

                            st.session_state.chapters["03_filled"] = filled_text
                            st.session_state.scores["03_filled"] = score_chapter(filled_text)
                            st.session_state.cache_stats.append({"label": "Fill", "info": fill_info})

                            fill_wc = len(filled_text.split())
                            st.info(f"Fill: {cut_wc} → {fill_wc} words ({fill_wc - cut_wc} added)")

    st.rerun()

elif not api_key:
    st.info("Enter your Anthropic API key in the sidebar to begin.")
